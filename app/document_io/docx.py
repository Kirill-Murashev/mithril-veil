from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.core.exceptions import EmptyExtractedText, MithrilVeilError


def read_docx_text(path: Path) -> str:
    """Extract paragraph and table cell text from a DOCX file."""
    try:
        document = Document(path)
    except (PackageNotFoundError, OSError, ValueError) as exc:
        raise MithrilVeilError("Cannot read DOCX file.") from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    combined = "\n".join(parts).strip()
    if not combined:
        raise EmptyExtractedText("DOCX file contains no extractable text.")
    return combined
