"""Generate synthetic DOCX/PDF/ODT fixtures in tests (not committed binaries)."""

from __future__ import annotations

import xml.sax.saxutils
import zipfile
from pathlib import Path

SYNTHETIC_CONTACT_LINE = "Контакт: test@example.local"
SYNTHETIC_INN = "7707083893"


def write_synthetic_rtf(path: Path, text: str = SYNTHETIC_CONTACT_LINE) -> None:
    """Write a minimal synthetic UTF-8 RTF fixture (plain-text extraction only)."""
    escaped = text.replace("\\", "\\\\").replace("{", "\\{").replace("}", "\\}")
    body = f"{{\\rtf1\\ansi {escaped}}}"
    path.write_text(body, encoding="utf-8")


def write_synthetic_rtf_cp1251(path: Path, text: str = SYNTHETIC_CONTACT_LINE) -> None:
    """Write a legacy cp1251 RTF with \\ansicpg1251 (synthetic Cyrillic-safe bytes only)."""
    payload = text.encode("cp1251")
    path.write_bytes(b"{\\rtf1\\ansi\\ansicpg1251 " + payload + b"}")


def write_synthetic_rtf_unicode_escapes(path: Path) -> None:
    """Write RTF using \\u escape sequences for Cyrillic (Иван + synthetic email)."""
    body = r"{\rtf1\ansi \u1048?\u1072?\u1085? test@example.local}"
    path.write_text(body, encoding="utf-8")


def write_synthetic_docx(path: Path, text: str = SYNTHETIC_CONTACT_LINE) -> None:
    from docx import Document

    document = Document()
    document.add_paragraph(text)
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "ООО Тестовая Организация"
    document.save(path)


def write_synthetic_pdf(path: Path, text: str = SYNTHETIC_CONTACT_LINE) -> None:
    from reportlab.pdfgen import canvas

    pdf = canvas.Canvas(str(path))
    pdf.drawString(72, 720, text)
    pdf.save()


def write_encrypted_pdf(path: Path, password: str = "synthetic-test") -> None:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.encrypt(password)
    with path.open("wb") as handle:
        writer.write(handle)


def write_synthetic_odt(
    path: Path,
    *,
    paragraphs: list[str] | None = None,
    table_rows: list[list[str]] | None = None,
    heading: str | None = None,
) -> None:
    """Build a minimal synthetic ODT (ZIP + content.xml) for plain-text extraction tests."""
    if paragraphs is None:
        paragraphs = [SYNTHETIC_CONTACT_LINE, f"ИНН: {SYNTHETIC_INN}"]

    parts = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"',
        ' xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0"',
        ' xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0">',
        "<office:body><office:text>",
    ]
    if heading:
        parts.append(f"<text:h>{xml.sax.saxutils.escape(heading)}</text:h>")
    for paragraph in paragraphs:
        parts.append(f"<text:p>{xml.sax.saxutils.escape(paragraph)}</text:p>")
    if table_rows:
        parts.append("<table:table>")
        for row in table_rows:
            parts.append("<table:table-row>")
            for cell in row:
                parts.append(
                    f"<table:table-cell><text:p>{xml.sax.saxutils.escape(cell)}</text:p>"
                    "</table:table-cell>"
                )
            parts.append("</table:table-row>")
        parts.append("</table:table>")
    parts.extend(["</office:text></office:body></office:document-content>"])
    content_xml = "".join(parts).encode("utf-8")

    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("mimetype", "application/vnd.oasis.opendocument.text")
        archive.writestr("content.xml", content_xml)


def write_blank_pdf(path: Path) -> None:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with path.open("wb") as handle:
        writer.write(handle)
